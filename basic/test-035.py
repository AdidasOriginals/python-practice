#!/usr/bin/env python
# _*_ coding: utf-8 _*_
# @Time : 2021/2/2 10:45 
# @Author : Edison
# @Version：V 0.1
# @File : test-035.py
# @desc :图像和办公文档处理
print('用Pillow操作图像'.center(50, '*'))
from PIL import Image

image = Image.open('image/dog.jpg')
image.format, image.size, image.mode
('JPEG', (500, 750), 'RGB')
# image.show()
print('剪裁图像'.center(50, '-'))
image = Image.open('image/dog.jpg')
rect = 80, 20, 310, 360
image.crop(rect)
# image.show()
print('生成缩略图'.center(50, '-'))
image1 = Image.open('image/dog.jpg')
size = 128, 128
image1.thumbnail(size)
# image1.show()
print('生成缩略图'.center(50, '-'))
image2 = Image.open('image/dog2.jpg')
image3 = Image.open('image/dog.jpg')
rect = 0, 0, 440, 418
dog_head = image3.crop(rect)
width, height = dog_head.size
image2.paste(dog_head.resize((int(width / 1.5), int(height / 1.5))), (172, 40))
# image2.show()
print('旋转和翻转'.center(50, '-'))
image4 = Image.open('image/dog.jpg')
image4.rotate(180)
# image4.show()
image4.transpose(Image.FLIP_LEFT_RIGHT)
# image4.show()
print('操作像素'.center(50, '-'))
image5 = Image.open('image/dog.jpg')
for x in range(80, 310):
    for y in range(20, 360):
        image5.putpixel((x, y), (128, 128, 128))
# image5.show()
print('滤镜效果'.center(50, '-'))
from PIL import Image, ImageFilter

image6 = Image.open('image/dog.jpg')
# image6.filter(ImageFilter.CONTOUR).show()


print('处理Excel电子表格'.center(50, '*'))
import datetime
from openpyxl import Workbook

wb = Workbook()
ws = wb.active
ws['A1'] = 42
ws.append([1, 2, 3])
ws['A3'] = datetime.datetime.now()
# wb.save('./data/test_save.xlsx')
print('处理Word文档'.center(50, '*'))
# pip3 install -i https://pypi.tuna.tsinghua.edu.cn/simple python-docx
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
from pathlib import Path
import base64

document = Document()
# 文档标题
document.add_heading('Document Title', 0)
# 段落
p = document.add_paragraph('A plain paragraph having some ')
# 加粗
p.add_run('bold').bold = True
p.add_run(' and some ')
# 斜体
p.add_run('italic.').italic = True
# 一级标题
document.add_heading('Heading, level 1', level=1)
# 段落样式
document.add_paragraph('Intense quote', style='Intense Quote')
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
# 添加图片
image_path = 'D:/PycharmProjects/python50/image/download/20210131092826291-6865224.jpg'
document.add_picture(image_path, width=Inches(1.25))
# 下面的不行
# with open('./image/download/20210131092826291-6865224.jpg','rb') as f:
#     document.add_picture(f.read(), width=Inches(1.25))


records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)
# 添加表格
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
# 分页符
document.add_page_break()
document.save('./data/demo.docx')
